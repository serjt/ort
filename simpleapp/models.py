# coding=utf-8
from __future__ import unicode_literals

import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
import xlsxwriter
from django.contrib.auth.models import User
from django.db import models

# Create your models here.
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

from diplomka import settings


def file_upload_to(instance, filename):
    return "%s" % filename


class Faculty(models.Model):
    class Meta:
        verbose_name = 'Отделение'
        verbose_name_plural = 'Отделения'

    name = models.CharField(max_length=100, verbose_name='Отделение')
    lessons = models.ManyToManyField('Lesson', verbose_name='Предметы')
    quota = models.IntegerField(default=0, verbose_name='Квота')
    filled_quota = models.IntegerField(default=0, verbose_name='Осталось')
    manager = models.ForeignKey(User, null=True, verbose_name='Админ')

    def __unicode__(self):
        return self.name

    def get_alumnis(self):
        now = datetime.datetime.now()
        tour = Tour.objects.filter(initial__lte=now, final__gte=now)[0]
        alumnis = Alumni.objects.filter(tour=tour, faculty=self)
        return alumnis


class Tour(models.Model):
    class Meta:
        verbose_name = 'Тур'
        verbose_name_plural = 'Туры'

    name = models.CharField(max_length=100, verbose_name='Название')
    initial = models.DateTimeField(null=True, blank=True, verbose_name='Начало')
    final = models.DateTimeField(null=True, blank=True, verbose_name='Конец')

    def __unicode__(self):
        return self.name


class Protocol(models.Model):
    class Meta:
        verbose_name = 'Протокол'
        verbose_name_plural = 'Протоколы'

    tour = models.ForeignKey(Tour, verbose_name='Тур')
    protocol = models.FileField(upload_to=file_upload_to, null=True, blank=True, verbose_name='Файл')
    date = models.DateField(auto_now=True, verbose_name='Дата')

    def save(self, *args, **kwargs):
        document = Document()

        obj_styles = document.styles
        obj_charstyle = obj_styles.add_style('style', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(12)
        obj_font.name = 'Times New Roman'
        dep = Faculty.objects.all()
        for i in dep:
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(u'« Кабыл алууга сунушталган абитуриенттердин тизмесин бекитүү тууралуу » ',
                      style='style').bold = True
            p1 = document.add_paragraph()
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(u'№ 3-Протокол', style='style').bold = True
            p1.add_run(u'«__» __ ____-жыл', style='style')
            p2 = document.add_paragraph().add_run(
                u'___-жылдын, __-июлундагы № 2-Протокол менен бекитилген '
                u'орундардын санына ылайык айрым категориялардын чектеринде абитуриенттерди конкурстук'
                u' тандоонун негизинде Гранттык комиссия Кыргыз-Түрк «Манас» университетине «%s» '
                u'адистиги боюнча абитуриенттерди кабыл алууга сунуштоо чечимин чыгарды:' % i.name, style='style')
            p3 = document.add_paragraph().add_run(u'- Бишкек ш. бүтүрүүчүлөрү (бөлүнгөн орундар) ', style='style')
            alumnis = Alumni.objects.filter(tour=self.tour, passed=True, faculty=i)
            table = document.add_table(rows=1, cols=4, style='Table Grid')
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = u'Иден №'
            hdr_cells[1].text = u'негизги тест'
            hdr_cells[2].text = u'кошумча тест'
            hdr_cells[3].text = u'суммасы'
            for j in alumnis.filter(place='Шаар'):
                row = table.add_row().cells
                row[0].text = j.ortId
                row[1].text = str(j.main)
                row[2].text = str(j.extra_num)
                row[3].text = str(j.summa)

            p3 = document.add_paragraph().add_run(u'- Област. борб. ж-а чакан шаар-дын бүтүрүүчүлөрү '
                                                  u'(бөлүнгөн орундар) ', style='style')
            table = document.add_table(rows=1, cols=4, style='Table Grid')
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = u'Иден №'
            hdr_cells[1].text = u'негизги тест'
            hdr_cells[2].text = u'кошумча тест'
            hdr_cells[3].text = u'суммасы'
            for j in alumnis.filter(place='Борбор'):
                row = table.add_row().cells
                row[0].text = j.ortId
                row[1].text = str(j.main)
                row[2].text = str(j.extra_num)
                row[3].text = str(j.summa)

            p3 = document.add_paragraph().add_run(u'- Айылдардын бүтүрүүчүлүрү (бөлүнгөн орундар)', style='style')
            table = document.add_table(rows=1, cols=4, style='Table Grid')
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = u'Иден №'
            hdr_cells[1].text = u'негизги тест'
            hdr_cells[2].text = u'кошумча тест'
            hdr_cells[3].text = u'суммасы'
            for j in alumnis.filter(place='Айыл'):
                row = table.add_row().cells
                row[0].text = j.ortId
                row[1].text = str(j.main)
                row[2].text = str(j.extra_num)
                row[3].text = str(j.summa)

            p3 = document.add_paragraph().add_run(u'- Бийик тоолуу райондордун бүтүрүүчүлөрү  (бөлүнгөн орундар) ',
                                                  style='style')
            table = document.add_table(rows=1, cols=4, style='Table Grid')
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = u'Иден №'
            hdr_cells[1].text = u'негизги тест'
            hdr_cells[2].text = u'кошумча тест'
            hdr_cells[3].text = u'суммасы'
            for j in alumnis.filter(place='Тоо'):
                row = table.add_row().cells
                row[0].text = j.ortId
                row[1].text = str(j.main)
                row[2].text = str(j.extra_num)
                row[3].text = str(j.summa)

            document.add_paragraph()
            p3 = document.add_paragraph().add_run('Гранттык комиссиянын төрагасы   '
                                                  '  _______________      А.А. Кулмырзаев ', style='style')
            document.add_page_break()
        document.save(settings.BASE_DIR + u'/static_in_env/media_root/protocol_%s.docx' % (self.tour.name))
        self.protocol = '/media/protocol_%s.docx' % (self.tour.name)
        super(Protocol, self).save()


class Otchet(models.Model):
    class Meta:
        verbose_name = 'Отчет'
        verbose_name_plural = 'Отчеты'

    tour = models.ForeignKey(Tour, null=True, verbose_name='Тур')
    department = models.ForeignKey(Faculty, null=True, verbose_name='Отделение')
    otchet = models.FileField(upload_to=file_upload_to, blank=True, null=True, verbose_name='Файл')
    date = models.DateField(auto_now=True, verbose_name='Дата')

    def save(self, *args, **kwargs):
        faculty = self.department
        tour = self.tour
        alumnis = Alumni.objects.filter(faculty=faculty, tour=tour, lgotnik__isnull=True)
        shaar = alumnis.filter(place=u'Шаар').exclude(olimpiadnik=True)
        borbor = alumnis.filter(place=u'Борбор').exclude(olimpiadnik=True)
        aiyl = alumnis.filter(place=u'Айыл').exclude(olimpiadnik=True)
        too = alumnis.filter(place=u'Тоо').exclude(olimpiadnik=True)
        olimpiadniki = alumnis.filter(olimpiadnik=True)
        name = settings.BASE_DIR + u'/static_in_env/media_root/otchet_%s_%s_%s.xlsx' % (
            tour.name, faculty.name, self.date)
        workbook = xlsxwriter.Workbook(name)
        worksheet = workbook.add_worksheet('Result')
        worksheet.set_column('A:A', 1.5)
        worksheet.set_column('B:Z', 5)
        format = workbook.add_format({'bg_color': 'red',
                                      'align': 'center',
                                      'valign': 'vcenter',
                                      'font_size': 7,
                                      'border': 1})
        worksheet.set_default_row(11)
        cell_format = workbook.add_format({'align': 'center',
                                           'valign': 'vcenter',
                                           'font_size': 7,
                                           'border': 1})
        cell_format_name = workbook.add_format({'align': 'center',
                                                'valign': 'vcenter',
                                                'font_size': 7,
                                                'border': 1})
        worksheet.merge_range('A1:K1', u"%s багыты боюнча" % faculty.name, cell_format)
        worksheet.merge_range('A2:K2', u"%sда катышкандардын тизмеси" % tour.name, cell_format)
        worksheet.merge_range('A3:K3', u"Кабыл алуу планы: %s" % faculty.filled_quota, cell_format)
        worksheet.merge_range('A5:A6', 'N', cell_format)
        worksheet.merge_range('B5:E5', u'Шаар', format)
        format_blue = workbook.add_format({'bg_color': 'blue',
                                           'align': 'center',
                                           'valign': 'vcenter',
                                           'font_size': 7,
                                           'border': 1})
        worksheet.merge_range('F5:I5', u'Кичи шаар жана обл. борборлор', format_blue)
        format_yellow = workbook.add_format({'bg_color': 'yellow',
                                             'align': 'center',
                                             'valign': 'vcenter',
                                             'font_size': 7,
                                             'border': 1})
        worksheet.merge_range('J5:M5', u'Айыл жергеси', format_yellow)
        format_purple = workbook.add_format({'bg_color': 'purple',
                                             'align': 'center',
                                             'valign': 'vcenter',
                                             'font_size': 7,
                                             'border': 1})
        worksheet.merge_range('N5:Q5', u'Бийик тоолу аймак', format_purple)
        format_white = workbook.add_format({'align': 'center',
                                            'valign': 'vcenter',
                                            'font_size': 7,
                                            'border': 1})
        worksheet.merge_range('R5:U5', u'Олимпиада жеңүүчүлөрү', format_white)
        worksheet.write('B6', u'Иден', cell_format_name)
        worksheet.write('C6', u'Негизги', cell_format_name)
        worksheet.write('D6', u'Кошумча', cell_format_name)
        worksheet.write('E6', u'Суммасы', cell_format_name)
        worksheet.write('F6', u'Иден', cell_format_name)
        worksheet.write('G6', u'Негизги', cell_format_name)
        worksheet.write('H6', u'Кошумча', cell_format_name)
        worksheet.write('I6', u'Суммасы', cell_format_name)
        worksheet.write('J6', u'Иден', cell_format_name)
        worksheet.write('K6', u'Негизги', cell_format_name)
        worksheet.write('L6', u'Кошумча', cell_format_name)
        worksheet.write('M6', u'Суммасы', cell_format_name)
        worksheet.write('N6', u'Иден', cell_format_name)
        worksheet.write('O6', u'Негизги', cell_format_name)
        worksheet.write('P6', u'Кошумча', cell_format_name)
        worksheet.write('Q6', u'Суммасы', cell_format_name)
        worksheet.write('R6', u'Иден', cell_format_name)
        worksheet.write('S6', u'Негизги', cell_format_name)
        worksheet.write('T6', u'Кошумча', cell_format_name)
        worksheet.write('U6', u'Суммасы', cell_format_name)
        worksheet.set_row(5, 30)
        l = [shaar.count(), aiyl.count(), too.count(), borbor.count(), olimpiadniki.count()]
        m = max(l)
        for i in 'ABCDEFGHIJKLMNOPQRSTU':
            for j in range(7, m + 7):
                worksheet.write('%s%s' % (i, j), None, cell_format_name)
        counter = 6
        for i in shaar:
            counter += 1
            worksheet.write('B%s' % str(counter), i.ortId, cell_format_name)
            worksheet.write('C%s' % str(counter), i.main, cell_format_name)
            worksheet.write('D%s' % str(counter), i.extra_num, cell_format_name)
            worksheet.write('E%s' % str(counter), i.summa, cell_format_name)
            worksheet.write('A%s' % str(counter), counter - 6, cell_format_name)
        counter = 6
        for i in borbor:
            counter += 1
            worksheet.write('F%s' % str(counter), i.ortId, cell_format_name)
            worksheet.write('G%s' % str(counter), i.main, cell_format_name)
            worksheet.write('H%s' % str(counter), i.extra_num, cell_format_name)
            worksheet.write('I%s' % str(counter), i.summa, cell_format_name)
            worksheet.write('A%s' % str(counter), counter - 6, cell_format_name)
        counter = 6
        for i in aiyl:
            counter += 1
            worksheet.write('J%s' % str(counter), i.ortId, cell_format_name)
            worksheet.write('K%s' % str(counter), i.main, cell_format_name)
            worksheet.write('L%s' % str(counter), i.extra_num, cell_format_name)
            worksheet.write('M%s' % str(counter), i.summa, cell_format_name)
            worksheet.write('A%s' % str(counter), counter - 6, cell_format_name)
        counter = 6
        for i in too:
            counter += 1
            worksheet.write('N%s' % str(counter), i.ortId, cell_format_name)
            worksheet.write('O%s' % str(counter), i.main, cell_format_name)
            worksheet.write('P%s' % str(counter), i.extra_num, cell_format_name)
            worksheet.write('Q%s' % str(counter), i.summa, cell_format_name)
            worksheet.write('A%s' % str(counter), counter - 6, cell_format_name)
        counter = 6
        for i in olimpiadniki:
            counter += 1
            worksheet.write('R%s' % str(counter), i.ortId, cell_format_name)
            worksheet.write('S%s' % str(counter), i.main, cell_format_name)
            worksheet.write('T%s' % str(counter), i.extra_num, cell_format_name)
            worksheet.write('U%s' % str(counter), i.summa, cell_format_name)
            worksheet.write('A%s' % str(counter), counter - 6, cell_format_name)
        m += 6
        worksheet.write('B%s' % (m + 3), u'Всего', cell_format_name)
        worksheet.write('C%s' % (m + 3), alumnis.count(), cell_format_name)
        worksheet.write('D%s' % (m + 5), u'Квота', cell_format_name)
        worksheet.merge_range('B%s:C%s' % (m + 5, m + 5), u'Шаар', cell_format_name)
        worksheet.write('G%s' % (m + 5), u'Квота', cell_format_name)
        worksheet.merge_range('E%s:F%s' % (m + 5, m + 5), u'Кичи шаар ж/а обл.', cell_format_name)
        worksheet.write('J%s' % (m + 5), u'Квота', cell_format_name)
        worksheet.merge_range('H%s:I%s' % (m + 5, m + 5), u'Айыл жергеси', cell_format_name)
        worksheet.write('M%s' % (m + 5), u'Квота', cell_format_name)
        worksheet.merge_range('K%s:L%s' % (m + 5, m + 5), u'Бийик тоолу айм.', cell_format_name)
        worksheet.write('P%s' % (m + 5), u'Квота', cell_format_name)
        worksheet.merge_range('N%s:O%s' % (m + 5, m + 5), u'Олимпиада жең.', cell_format_name)
        n = 0
        k = 0
        o = 0
        p = 0
        if alumnis.count() != 0:
            if shaar.count() != 0:
                n = shaar.count() * (faculty.filled_quota - olimpiadniki.count()) / float(
                    alumnis.count() - olimpiadniki.count())
            if borbor.count() != 0:
                k = borbor.count() * (faculty.filled_quota - olimpiadniki.count()) / float(
                    alumnis.count() - olimpiadniki.count())
            if aiyl.count() != 0:
                o = aiyl.count() * (faculty.filled_quota - olimpiadniki.count()) / float(
                    alumnis.count() - olimpiadniki.count())
            if too.count() != 0:
                p = too.count() * (faculty.filled_quota - olimpiadniki.count()) / float(
                    alumnis.count() - olimpiadniki.count())
        worksheet.write('D%s' % (m + 6), n, cell_format_name)
        worksheet.merge_range('B%s:C%s' % (m + 6, m + 6), shaar.count(), cell_format_name)
        worksheet.write('G%s' % (m + 6), k, cell_format_name)
        worksheet.merge_range('E%s:F%s' % (m + 6, m + 6), borbor.count(), cell_format_name)
        worksheet.write('J%s' % (m + 6), o, cell_format_name)
        worksheet.merge_range('H%s:I%s' % (m + 6, m + 6), aiyl.count(), cell_format_name)
        worksheet.write('M%s' % (m + 6), p, cell_format_name)
        worksheet.merge_range('K%s:L%s' % (m + 6, m + 6), too.count(), cell_format_name)
        worksheet.write('P%s' % (m + 6), olimpiadniki.count(), cell_format_name)
        worksheet.merge_range('N%s:O%s' % (m + 6, m + 6), olimpiadniki.count(), cell_format_name)
        barcode_worksheet = workbook.add_worksheet('Barcode')
        barcode_worksheet.set_column('A:A', 1.5)
        barcode_worksheet.set_column('B:B', 28)
        barcode_worksheet.set_column('C:C', 20)
        barcode_worksheet.set_column('D:D', 18)
        barcode_worksheet.set_column('E:E', 5)
        barcode_worksheet.set_column('F:F', 14)
        barcode_worksheet.set_column('G:G', 12)
        barcode_format = workbook.add_format({'align': 'center',
                                              'valign': 'vcenter',
                                              'font_size': 10,
                                              'border': 1})
        g = ord('G')
        for i in faculty.lessons.all():
            g += 1
            barcode_worksheet.write('%s1' % chr(g), i.name, barcode_format)
        barcode_worksheet.write('A1', u'№', barcode_format)
        barcode_worksheet.write('B1', u'поле для ввода со сканера ШК', barcode_format)
        barcode_worksheet.write('C1', u'регистрационный номер', barcode_format)
        barcode_worksheet.write('D1', u'цвет сертификата', barcode_format)
        barcode_worksheet.write('E1', u'тур', barcode_format)
        barcode_worksheet.write('F1', u'красный аттестат', barcode_format)
        barcode_worksheet.write('G1', u'телефон', barcode_format)
        count = 1
        for i in alumnis.all():
            count += 1
            barcode_worksheet.write('B%s' % count, i.barcode, barcode_format)
            barcode_worksheet.write('A%s' % count, count - 1, barcode_format)
            barcode_worksheet.write('C%s' % count, i.ortId, barcode_format)
            if i.place == u'Шаар':
                barcode_worksheet.write('D%s' % count, u'К', format)
            elif i.place == u'Борбор':
                barcode_worksheet.write('D%s' % count, u'С', format_blue)
            elif i.place == u'Айыл':
                barcode_worksheet.write('D%s' % count, u'Ж', format_yellow)
            else:
                barcode_worksheet.write('D%s' % count, u'Ф', format_purple)
            barcode_worksheet.write('E%s' % count, i.tour.name, barcode_format)
            if i.atestat:
                barcode_worksheet.write('F%s' % count, '*', barcode_format)
            else:
                barcode_worksheet.write('F%s' % count, ' ', barcode_format)
            barcode_worksheet.write('G%s' % count, i.phone, barcode_format)
            g = ord('G')
            for j in faculty.lessons.all():
                g += 1
                barcode_worksheet.write('%s%s' % (chr(g), count), AlumniLesson.objects.get(lesson=j, alumni=i).grade,
                                        barcode_format)

        journal_worksheet = workbook.add_worksheet('Journal')
        journal_worksheet.set_column('A:A', 1.5)
        journal_worksheet.set_column('B:G', 10)
        journal_format = workbook.add_format({'align': 'center',
                                              'valign': 'vcenter',
                                              'font_size': 8,
                                              'border': 1})
        journal_worksheet.write('A1', u'№', journal_format)
        journal_worksheet.write('B1', u'Идент.', journal_format)
        journal_worksheet.write('C1', u'Основной', journal_format)
        journal_worksheet.write('D1', u'Доп', journal_format)
        journal_worksheet.write('E1', u'Сумма', journal_format)
        journal_worksheet.write('F1', u'Категория', journal_format)
        journal_worksheet.write('G1', u'Атестат', journal_format)
        journal_worksheet.set_default_row(11)
        c = 0
        for i in alumnis.all():
            c += 1
            journal_worksheet.write('A%s' % (c + 1), c, journal_format)
            journal_worksheet.write('B%s' % (c + 1), i.ortId, journal_format)
            journal_worksheet.write('C%s' % (c + 1), i.main, journal_format)
            journal_worksheet.write('D%s' % (c + 1), i.extra_num, journal_format)
            journal_worksheet.write('E%s' % (c + 1), i.summa, journal_format)
            if i.place == u'Шаар':
                journal_worksheet.write('F%s' % (c + 1), u'К', format)
            elif i.place == u'Борбор':
                journal_worksheet.write('F%s' % (c + 1), u'С', format_blue)
            elif i.place == u'Айыл':
                journal_worksheet.write('F%s' % (c + 1), u'Ж', format_yellow)
            else:
                journal_worksheet.write('F%s' % (c + 1), u'Ф', format_purple)
            if i.atestat:
                journal_worksheet.write('G%s' % (c + 1), '*', journal_format)
            else:
                journal_worksheet.write('G%s' % (c + 1), ' ', journal_format)
        workbook.close()
        self.otchet = '/media/otchet_%s_%s_%s.xlsx' % (tour.name, faculty.name, self.date)
        super(Otchet, self).save()


class Lesson(models.Model):
    class Meta:
        verbose_name = 'предмет'
        verbose_name_plural = 'предметы'

    name = models.CharField(max_length=100, verbose_name='Название')

    def __unicode__(self):
        return self.name


class Lgotnik(models.Model):
    class Meta:
        verbose_name = 'Льготник'
        verbose_name_plural = 'Льготники'

    name = models.CharField(max_length=100, verbose_name='Название')
    quota = models.IntegerField(default=0, verbose_name='Квота')
    filled_quota = models.IntegerField(default=0, verbose_name='Осталось')
    date = models.DateField(auto_now=True, verbose_name='Дата')

    def __unicode__(self):
        return self.name


class Alumni(models.Model):
    class Meta:
        ordering = ['-date']
        verbose_name = 'абитуриента'
        verbose_name_plural = 'абитуриенты'

    choices = (
        ('Шаар', 'Шаар'),
        ('Борбор', 'Борбор'),
        ('Айыл', 'Айыл'),
        ('Тоо', 'Тоо'),
    )
    barcode = models.CharField(max_length=100, null=True)
    ortId = models.CharField(max_length=100, verbose_name='ID')
    tour = models.ForeignKey(Tour, null=True, verbose_name='Тур')
    faculty = models.ForeignKey(Faculty, null=True, verbose_name='Отделение')
    place = models.CharField(max_length=100, null=True, choices=choices, verbose_name='Местность')
    extra_num = models.IntegerField(default=0, verbose_name='Доп.')
    main = models.IntegerField(default=0, verbose_name='Осн.')
    atestat = models.BooleanField(default=False, verbose_name='Красный аттестат')
    lgotnik = models.ForeignKey(Lgotnik, blank=True, null=True, verbose_name='Льготник')
    olimpiadnik = models.BooleanField(default=False, verbose_name='Олимпиадник')
    passed = models.BooleanField(default=False)
    summa = models.IntegerField(default=0, verbose_name='Сумма')
    phone = models.CharField(max_length=100, null=True, verbose_name='Номер')
    date = models.DateTimeField(auto_now=True, verbose_name='Дата')

    # def __unicode__(self):
    #     return self.barcode

    def get_main(self):
        less = Lesson.objects.get(name='Основной')
        alumnilesson = AlumniLesson.objects.get(lesson=less, alumni=self)
        return alumnilesson.grade


class AlumniLesson(models.Model):
    alumni = models.ForeignKey(Alumni, null=True, verbose_name='абитуриент')
    lesson = models.ForeignKey(Lesson, null=True, verbose_name='предмет')
    grade = models.IntegerField(default=0, verbose_name='Балл')

    def __unicode__(self):
        return self.lesson.name


class OtchetLgotnik(models.Model):
    class Meta:
        verbose_name = 'Отчет(Льготник)'
        verbose_name_plural = 'Отчет(Льготники)'
    tour = models.ForeignKey(Tour, verbose_name="Тур")
    lgotnik = models.ForeignKey(Lgotnik, verbose_name='Льготник')
    file = models.FileField(upload_to=file_upload_to, blank=True, null=True, verbose_name='Файл')
    date = models.DateField(auto_now=True, verbose_name='Дата')

    def save(self, *args, **kwargs):
        tour = self.tour
        lgotnik = self.lgotnik
        alumnis = Alumni.objects.filter(tour=tour, lgotnik=lgotnik)
        shaar = alumnis.filter(place=u'Шаар').exclude(olimpiadnik=True)
        borbor = alumnis.filter(place=u'Борбор').exclude(olimpiadnik=True)
        aiyl = alumnis.filter(place=u'Айыл').exclude(olimpiadnik=True)
        too = alumnis.filter(place=u'Тоо').exclude(olimpiadnik=True)
        olimpiadniki = alumnis.filter(olimpiadnik=True)
        name = settings.BASE_DIR + u'/static_in_env/media_root/otchet_%s_%s_%s.xlsx' % (
            tour.name, lgotnik.name, self.date)
        workbook = xlsxwriter.Workbook(name)
        worksheet = workbook.add_worksheet('Result')
        worksheet.set_column('A:A', 1.5)
        worksheet.set_column('B:Z', 5)
        format = workbook.add_format({'bg_color': 'red',
                                      'align': 'center',
                                      'valign': 'vcenter',
                                      'font_size': 7,
                                      'border': 1})
        worksheet.set_default_row(11)
        cell_format = workbook.add_format({'align': 'center',
                                           'valign': 'vcenter',
                                           'font_size': 7,
                                           'border': 1})
        cell_format_name = workbook.add_format({'align': 'center',
                                                'valign': 'vcenter',
                                                'font_size': 7,
                                                'border': 1})
        worksheet.merge_range('A1:K1', u"%s багыты боюнча" % lgotnik.name, cell_format)
        worksheet.merge_range('A2:K2', u"%sда катышкандардын тизмеси" % tour.name, cell_format)
        worksheet.merge_range('A3:K3', u"Кабыл алуу планы: %s" % lgotnik.filled_quota, cell_format)
        worksheet.merge_range('A5:A6', 'N', cell_format)
        worksheet.merge_range('B5:E5', u'Шаар', format)
        format_blue = workbook.add_format({'bg_color': 'blue',
                                           'align': 'center',
                                           'valign': 'vcenter',
                                           'font_size': 7,
                                           'border': 1})
        worksheet.merge_range('F5:I5', u'Кичи шаар жана обл. борборлор', format_blue)
        format_yellow = workbook.add_format({'bg_color': 'yellow',
                                             'align': 'center',
                                             'valign': 'vcenter',
                                             'font_size': 7,
                                             'border': 1})
        worksheet.merge_range('J5:M5', u'Айыл жергеси', format_yellow)
        format_purple = workbook.add_format({'bg_color': 'purple',
                                             'align': 'center',
                                             'valign': 'vcenter',
                                             'font_size': 7,
                                             'border': 1})
        worksheet.merge_range('N5:Q5', u'Бийик тоолу аймак', format_purple)
        format_white = workbook.add_format({'align': 'center',
                                            'valign': 'vcenter',
                                            'font_size': 7,
                                            'border': 1})
        worksheet.merge_range('R5:U5', u'Олимпиада жеңүүчүлөрү', format_white)
        worksheet.write('B6', u'Иден', cell_format_name)
        worksheet.write('C6', u'Негизги', cell_format_name)
        worksheet.write('D6', u'Кошумча', cell_format_name)
        worksheet.write('E6', u'Суммасы', cell_format_name)
        worksheet.write('F6', u'Иден', cell_format_name)
        worksheet.write('G6', u'Негизги', cell_format_name)
        worksheet.write('H6', u'Кошумча', cell_format_name)
        worksheet.write('I6', u'Суммасы', cell_format_name)
        worksheet.write('J6', u'Иден', cell_format_name)
        worksheet.write('K6', u'Негизги', cell_format_name)
        worksheet.write('L6', u'Кошумча', cell_format_name)
        worksheet.write('M6', u'Суммасы', cell_format_name)
        worksheet.write('N6', u'Иден', cell_format_name)
        worksheet.write('O6', u'Негизги', cell_format_name)
        worksheet.write('P6', u'Кошумча', cell_format_name)
        worksheet.write('Q6', u'Суммасы', cell_format_name)
        worksheet.write('R6', u'Иден', cell_format_name)
        worksheet.write('S6', u'Негизги', cell_format_name)
        worksheet.write('T6', u'Кошумча', cell_format_name)
        worksheet.write('U6', u'Суммасы', cell_format_name)
        worksheet.set_row(5, 30)
        l = [shaar.count(), aiyl.count(), too.count(), borbor.count(), olimpiadniki.count()]
        m = max(l)
        for i in 'ABCDEFGHIJKLMNOPQRSTU':
            for j in range(7, m + 7):
                worksheet.write('%s%s' % (i, j), None, cell_format_name)
        counter = 6
        for i in shaar:
            counter += 1
            worksheet.write('B%s' % str(counter), i.ortId, cell_format_name)
            worksheet.write('C%s' % str(counter), i.main, cell_format_name)
            worksheet.write('D%s' % str(counter), i.extra_num, cell_format_name)
            worksheet.write('E%s' % str(counter), i.summa, cell_format_name)
            worksheet.write('A%s' % str(counter), counter - 6, cell_format_name)
        counter = 6
        for i in borbor:
            counter += 1
            worksheet.write('F%s' % str(counter), i.ortId, cell_format_name)
            worksheet.write('G%s' % str(counter), i.main, cell_format_name)
            worksheet.write('H%s' % str(counter), i.extra_num, cell_format_name)
            worksheet.write('I%s' % str(counter), i.summa, cell_format_name)
            worksheet.write('A%s' % str(counter), counter - 6, cell_format_name)
        counter = 6
        for i in aiyl:
            counter += 1
            worksheet.write('J%s' % str(counter), i.ortId, cell_format_name)
            worksheet.write('K%s' % str(counter), i.main, cell_format_name)
            worksheet.write('L%s' % str(counter), i.extra_num, cell_format_name)
            worksheet.write('M%s' % str(counter), i.summa, cell_format_name)
            worksheet.write('A%s' % str(counter), counter - 6, cell_format_name)
        counter = 6
        for i in too:
            counter += 1
            worksheet.write('N%s' % str(counter), i.ortId, cell_format_name)
            worksheet.write('O%s' % str(counter), i.main, cell_format_name)
            worksheet.write('P%s' % str(counter), i.extra_num, cell_format_name)
            worksheet.write('Q%s' % str(counter), i.summa, cell_format_name)
            worksheet.write('A%s' % str(counter), counter - 6, cell_format_name)
        counter = 6
        for i in olimpiadniki:
            counter += 1
            worksheet.write('R%s' % str(counter), i.ortId, cell_format_name)
            worksheet.write('S%s' % str(counter), i.main, cell_format_name)
            worksheet.write('T%s' % str(counter), i.extra_num, cell_format_name)
            worksheet.write('U%s' % str(counter), i.summa, cell_format_name)
            worksheet.write('A%s' % str(counter), counter - 6, cell_format_name)
        m += 6
        worksheet.write('B%s' % (m + 3), u'Всего', cell_format_name)
        worksheet.write('C%s' % (m + 3), alumnis.count(), cell_format_name)
        worksheet.write('D%s' % (m + 5), u'Квота', cell_format_name)
        worksheet.merge_range('B%s:C%s' % (m + 5, m + 5), u'Шаар', cell_format_name)
        worksheet.write('G%s' % (m + 5), u'Квота', cell_format_name)
        worksheet.merge_range('E%s:F%s' % (m + 5, m + 5), u'Кичи шаар ж/а обл.', cell_format_name)
        worksheet.write('J%s' % (m + 5), u'Квота', cell_format_name)
        worksheet.merge_range('H%s:I%s' % (m + 5, m + 5), u'Айыл жергеси', cell_format_name)
        worksheet.write('M%s' % (m + 5), u'Квота', cell_format_name)
        worksheet.merge_range('K%s:L%s' % (m + 5, m + 5), u'Бийик тоолу айм.', cell_format_name)
        worksheet.write('P%s' % (m + 5), u'Квота', cell_format_name)
        worksheet.merge_range('N%s:O%s' % (m + 5, m + 5), u'Олимпиада жең.', cell_format_name)
        n = 0
        k = 0
        o = 0
        p = 0
        if alumnis.count() != 0:
            if shaar.count() != 0:
                n = shaar.count() * (lgotnik.filled_quota - olimpiadniki.count()) / float(
                    alumnis.count() - olimpiadniki.count())
            if borbor.count() != 0:
                k = borbor.count() * (lgotnik.filled_quota - olimpiadniki.count()) / float(
                    alumnis.count() - olimpiadniki.count())
            if aiyl.count() != 0:
                o = aiyl.count() * (lgotnik.filled_quota - olimpiadniki.count()) / float(
                    alumnis.count() - olimpiadniki.count())
            if too.count() != 0:
                p = too.count() * (lgotnik.filled_quota - olimpiadniki.count()) / float(
                    alumnis.count() - olimpiadniki.count())
        worksheet.write('D%s' % (m + 6), n, cell_format_name)
        worksheet.merge_range('B%s:C%s' % (m + 6, m + 6), shaar.count(), cell_format_name)
        worksheet.write('G%s' % (m + 6), k, cell_format_name)
        worksheet.merge_range('E%s:F%s' % (m + 6, m + 6), borbor.count(), cell_format_name)
        worksheet.write('J%s' % (m + 6), o, cell_format_name)
        worksheet.merge_range('H%s:I%s' % (m + 6, m + 6), aiyl.count(), cell_format_name)
        worksheet.write('M%s' % (m + 6), p, cell_format_name)
        worksheet.merge_range('K%s:L%s' % (m + 6, m + 6), too.count(), cell_format_name)
        worksheet.write('P%s' % (m + 6), olimpiadniki.count(), cell_format_name)
        worksheet.merge_range('N%s:O%s' % (m + 6, m + 6), olimpiadniki.count(), cell_format_name)
        barcode_worksheet = workbook.add_worksheet('Barcode')
        barcode_worksheet.set_column('A:A', 1.5)
        barcode_worksheet.set_column('B:B', 28)
        barcode_worksheet.set_column('C:C', 20)
        barcode_worksheet.set_column('D:D', 18)
        barcode_worksheet.set_column('E:E', 5)
        barcode_worksheet.set_column('F:F', 14)
        barcode_worksheet.set_column('G:G', 12)
        barcode_format = workbook.add_format({'align': 'center',
                                              'valign': 'vcenter',
                                              'font_size': 10,
                                              'border': 1})
        g = ord('G')
        for i in Lesson.objects.all():
            g += 1
            barcode_worksheet.write('%s1' % chr(g), i.name, barcode_format)
        barcode_worksheet.write('A1', u'№', barcode_format)
        barcode_worksheet.write('B1', u'поле для ввода со сканера ШК', barcode_format)
        barcode_worksheet.write('C1', u'регистрационный номер', barcode_format)
        barcode_worksheet.write('D1', u'цвет сертификата', barcode_format)
        barcode_worksheet.write('E1', u'тур', barcode_format)
        barcode_worksheet.write('F1', u'отделение', barcode_format)
        barcode_worksheet.write('G1', u'телефон', barcode_format)
        count = 1
        for i in alumnis.all():
            count += 1
            barcode_worksheet.write('B%s' % count, i.barcode, barcode_format)
            barcode_worksheet.write('A%s' % count, count - 1, barcode_format)
            barcode_worksheet.write('C%s' % count, i.ortId, barcode_format)
            if i.place == u'Шаар':
                barcode_worksheet.write('D%s' % count, u'К', format)
            elif i.place == u'Борбор':
                barcode_worksheet.write('D%s' % count, u'С', format_blue)
            elif i.place == u'Айыл':
                barcode_worksheet.write('D%s' % count, u'Ж', format_yellow)
            else:
                barcode_worksheet.write('D%s' % count, u'Ф', format_purple)
            barcode_worksheet.write('E%s' % count, i.tour.name, barcode_format)
            barcode_worksheet.write('F%s' % count, i.faculty.name, barcode_format)
            barcode_worksheet.write('G%s' % count, i.phone, barcode_format)
            g = ord('G')
            for j in Lesson.objects.all():
                g += 1
                try:
                    barcode_worksheet.write('%s%s' % (chr(g), count), AlumniLesson.objects.get(lesson=j, alumni=i).grade,
                                        barcode_format)
                except:
                    barcode_worksheet.write('%s%s' % (chr(g), count), '000',
                                        barcode_format)

        journal_worksheet = workbook.add_worksheet('Journal')
        journal_worksheet.set_column('A:A', 1.5)
        journal_worksheet.set_column('B:G', 10)
        journal_format = workbook.add_format({'align': 'center',
                                              'valign': 'vcenter',
                                              'font_size': 8,
                                              'border': 1})
        journal_worksheet.write('A1', u'№', journal_format)
        journal_worksheet.write('B1', u'Идент.', journal_format)
        journal_worksheet.write('C1', u'Основной', journal_format)
        journal_worksheet.write('D1', u'Доп', journal_format)
        journal_worksheet.write('E1', u'Сумма', journal_format)
        journal_worksheet.write('F1', u'Категория', journal_format)
        journal_worksheet.write('G1', u'Атестат', journal_format)
        journal_worksheet.set_default_row(11)
        c = 0
        for i in alumnis.all():
            c += 1
            journal_worksheet.write('A%s' % (c + 1), c, journal_format)
            journal_worksheet.write('B%s' % (c + 1), i.ortId, journal_format)
            journal_worksheet.write('C%s' % (c + 1), i.main, journal_format)
            journal_worksheet.write('D%s' % (c + 1), i.extra_num, journal_format)
            journal_worksheet.write('E%s' % (c + 1), i.summa, journal_format)
            if i.place == u'Шаар':
                journal_worksheet.write('F%s' % (c + 1), u'К', format)
            elif i.place == u'Борбор':
                journal_worksheet.write('F%s' % (c + 1), u'С', format_blue)
            elif i.place == u'Айыл':
                journal_worksheet.write('F%s' % (c + 1), u'Ж', format_yellow)
            else:
                journal_worksheet.write('F%s' % (c + 1), u'Ф', format_purple)
            if i.atestat:
                journal_worksheet.write('G%s' % (c + 1), '*', journal_format)
            else:
                journal_worksheet.write('G%s' % (c + 1), ' ', journal_format)
        workbook.close()
        self.file = '/media/otchet_%s_%s_%s.xlsx' % (tour.name, lgotnik.name, self.date)
        super(OtchetLgotnik, self).save()
